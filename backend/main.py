import os
import uuid
import asyncio
from db import engine
from sql_models import Base


from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse
from starlette.middleware.sessions import SessionMiddleware
from authlib.integrations.starlette_client import OAuth

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from auth import get_current_user, create_access_token, hash_password, verify_password
from models import (
    ResearchRequest,
    UserRegister,
    UserLogin,
    ChangePasswordRequest,
    AgentBuilderRequest,
    RunAgentRequest,
)
from services.llm_router import ask_llm
from services.smart_router import classify_intent
from services.file_service import extract_file_content
from sql_database import (
    save_research,
    get_user_history,
    delete_research,
    create_user,
    authenticate_user,
)
from db import SessionLocal
from sql_models import User, UploadedDocument, GeneratedAgent, MarketplaceAgent


app = FastAPI(
    title="NeuroFlow AI",
    version="1.0.0"
)

Base.metadata.create_all(bind=engine)


app.add_middleware(
    SessionMiddleware,
    secret_key=os.getenv("JWT_SECRET_KEY", "neuroflow-session-secret")
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

oauth = OAuth()

oauth.register(
    name="google",
    client_id=os.getenv("GOOGLE_CLIENT_ID"),
    client_secret=os.getenv("GOOGLE_CLIENT_SECRET"),
    server_metadata_url="https://accounts.google.com/.well-known/openid-configuration",
    client_kwargs={"scope": "openid email profile"}
)

UPLOAD_DIR = "uploads"
DOWNLOAD_DIR = "downloads"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DOWNLOAD_DIR, exist_ok=True)


def get_selected_model(request: Request):
    return request.headers.get("X-AI-Model", "auto")


@app.get("/")
async def root():
    return {"message": "NeuroFlow AI Backend Running"}


@app.get("/health")
async def health():
    return {"status": "healthy", "service": "NeuroFlow AI"}


@app.post("/register")
async def register(user: UserRegister):
    success = create_user(user.email, user.password)

    if not success:
        raise HTTPException(status_code=400, detail="User already exists")

    return {"message": "Registration successful"}


@app.post("/login")
async def login(user: UserLogin):
    db_user = authenticate_user(user.email, user.password)

    if not db_user:
        raise HTTPException(status_code=401, detail="Invalid credentials")

    token = create_access_token({"sub": db_user["email"]})

    return {
        "access_token": token,
        "token_type": "bearer"
    }


@app.get("/me")
async def me(user=Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    return {"email": user}


@app.post("/research")
async def research(
    request: ResearchRequest,
    http_request: Request,
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(
            status_code=401,
            detail="Unauthorized"
        )

    selected_model = get_selected_model(http_request)
    topic = request.topic.strip()

    if not topic:
        raise HTTPException(
            status_code=400,
            detail="Topic is required"
        )

    print("REQUEST RECEIVED")
    print("USER:", user)
    print("MODEL:", selected_model)
    print("TOPIC:", topic)

    try:
        intent = await asyncio.wait_for(
            classify_intent(topic),
            timeout=20
        )
    except asyncio.TimeoutError:
        intent = "RESEARCH"
    except Exception as e:
        print("INTENT ERROR:", e)
        intent = "RESEARCH"

    print("INTENT:", intent)

    if intent == "CHAT":
        prompt = f"""
You are NeuroFlow AI chat assistant.

User Message:
{topic}

Rules:
- Give a clear and helpful answer.
- Keep answer simple.
- Reply in the same language as user if possible.
"""

        print("CHAT LLM START")

        try:
            answer = await asyncio.wait_for(
                ask_llm(prompt, selected_model),
                timeout=60
            )
        except asyncio.TimeoutError:
            raise HTTPException(
                status_code=504,
                detail="AI model timeout. Please select Gemini or Groq and try again."
            )
        except Exception as e:
            print("CHAT LLM ERROR:", e)
            raise HTTPException(
                status_code=500,
                detail=f"AI error: {str(e)}"
            )

        print("CHAT LLM DONE")

        saved = save_research(
            user_email=user,
            topic=topic,
            report=answer,
            logs=[
                "Chat Mode Activated",
                f"AI Model Used: {selected_model}"
            ]
        )

        return {
            "id": saved["id"],
            "topic": topic,
            "answer": answer,
            "report": answer,
            "logs": [
                "Chat Mode Activated",
                f"AI Model Used: {selected_model}"
            ]
        }

    prompt = f"""
You are NeuroFlow AI Research Assistant.

Selected AI Model:
{selected_model}

Research Topic:
{topic}

Create a professional, structured research report.

Required format:
# Executive Summary
# Key Findings
# Technical Analysis
# Comparison / Summary Table
# Risks and Challenges
# Recommendations
# Conclusion
# Sources

Rules:
- Use markdown.
- Use headings.
- Use bullet points.
- Use table where useful.
- Keep language professional and presentation-ready.
"""

    print("RESEARCH LLM START")

    try:
        answer = await asyncio.wait_for(
            ask_llm(prompt, selected_model),
            timeout=90
        )
    except asyncio.TimeoutError:
        raise HTTPException(
            status_code=504,
            detail="AI model timeout. Please select Gemini or Groq and try again."
        )
    except Exception as e:
        print("RESEARCH LLM ERROR:", e)
        raise HTTPException(
            status_code=500,
            detail=f"AI error: {str(e)}"
        )

    print("RESEARCH LLM DONE")

    saved = save_research(
        user_email=user,
        topic=topic,
        report=answer,
        logs=[
            "Research Mode Activated",
            f"AI Model Used: {selected_model}"
        ]
    )

    return {
        "id": saved["id"],
        "topic": topic,
        "answer": answer,
        "report": answer,
        "logs": [
            "Research Mode Activated",
            f"AI Model Used: {selected_model}"
        ]
    }

@app.get("/history")
async def history(user=Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    return get_user_history(user)


@app.delete("/history/{research_id}")
async def delete_history(
    research_id: str,
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    success = delete_research(research_id, user)

    if not success:
        raise HTTPException(status_code=404, detail="Research not found")

    return {"message": "Deleted successfully"}


@app.post("/upload-file")
async def upload_file(
    file: UploadFile = File(...),
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    allowed_extensions = [
        "pdf", "docx", "txt", "csv", "xlsx", "xls",
        "pptx", "png", "jpg", "jpeg", "json"
    ]

    original_filename = file.filename or "uploaded_file"
    extension = original_filename.split(".")[-1].lower()

    if extension not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    safe_filename = f"{uuid.uuid4()}.{extension}"
    filepath = os.path.join(UPLOAD_DIR, safe_filename)

    with open(filepath, "wb") as f:
        f.write(await file.read())

    content, file_type = await extract_file_content(filepath, extension)

    if not content:
        raise HTTPException(
            status_code=400,
            detail="Could not extract content from file"
        )

    db = SessionLocal()

    try:
        document = UploadedDocument(
            user_email=user,
            filename=original_filename,
            file_type=file_type,
            content=content
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return {
            "message": "File uploaded successfully",
            "document_id": document.id,
            "filename": original_filename,
            "file_type": file_type,
            "characters": len(content),
            "preview": content[:500]
        }

    finally:
        db.close()


@app.post("/ask-document")
async def ask_document(
    request: Request,
    data: dict,
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    selected_model = get_selected_model(request)

    document_id = data.get("document_id")
    question = data.get("question", "").strip()

    if not document_id:
        raise HTTPException(status_code=400, detail="document_id is required")

    if not question:
        raise HTTPException(status_code=400, detail="question is required")

    db = SessionLocal()

    try:
        document = (
            db.query(UploadedDocument)
            .filter(
                UploadedDocument.id == document_id,
                UploadedDocument.user_email == user
            )
            .first()
        )

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        prompt = f"""
Answer the question using only the uploaded file content.

Selected AI Model:
{selected_model}

File Name:
{document.filename}

File Type:
{document.file_type}

File Content:
{document.content[:12000]}

Question:
{question}

Rules:
- Give a clear answer.
- If information is not present, say "Information not found in uploaded file."
"""

        answer = await ask_llm(prompt, selected_model)

        saved = save_research(
            user_email=user,
            topic=f"Document Q&A: {document.filename}",
            report=answer,
            logs=[
                "Document Q&A Mode Activated",
                f"Answered from {document.filename}",
                f"AI Model Used: {selected_model}"
            ]
        )

        return {
            "id": saved["id"],
            "answer": answer,
            "report": answer,
            "logs": [
                "Document Q&A Mode Activated",
                f"Answered from {document.filename}",
                f"AI Model Used: {selected_model}"
            ]
        }

    finally:
        db.close()


@app.get("/dashboard-stats")
async def dashboard_stats(user=Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    db = SessionLocal()

    try:
        history_data = get_user_history(user)

        documents = (
            db.query(UploadedDocument)
            .filter(UploadedDocument.user_email == user)
            .all()
        )

        agents_created = len([
            item for item in history_data
            if "Agent Builder" in item.get("topic", "")
        ])

        recent_activity = []

        for item in history_data[:5]:
            recent_activity.append({
                "type": "research",
                "title": item.get("topic", "Research"),
                "description": item.get("report", "")[:80]
            })

        for doc in documents[:5]:
            recent_activity.append({
                "type": doc.file_type,
                "title": doc.filename,
                "description": f"{doc.file_type.upper()} uploaded"
            })

        return {
            "total_research": len(history_data),
            "documents": len(documents),
            "agents_created": agents_created,
            "hours_saved": len(history_data) * 2,
            "recent_activity": recent_activity[:6]
        }

    finally:
        db.close()


@app.post("/build-agent")
async def build_agent(
    request: AgentBuilderRequest,
    http_request: Request,
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    selected_model = get_selected_model(http_request)
    agent_type = request.agent_type.strip()

    if len(agent_type) < 3:
        raise HTTPException(
            status_code=400,
            detail="Please enter a valid agent type."
        )

    prompt = f"""
You are an expert AI Agent Architect.

Selected AI Model:
{selected_model}

Create a ready-to-use professional AI agent.

User Requirement:
{agent_type}

Return output in this exact format:

# Agent Name
# Purpose
# Role
# Goal
# Backstory
# Core Responsibilities
# Tools Required
# Workflow
# Input Format
# Output Format
# Ready-to-Use System Prompt
# FastAPI Endpoint Example
# Python Agent Class Code

Rules:
- Make it practical and professional.
- Make it ready to use in a Python/FastAPI project.
- Do not give generic content.
"""

    answer = await ask_llm(prompt, selected_model)

    db = SessionLocal()

    try:
        generated_agent = GeneratedAgent(
            user_email=user,
            agent_name=agent_type[:60],
            agent_type=agent_type,
            role="AI Domain Agent",
            goal=f"Create a ready-to-use AI agent for {agent_type}",
            prompt=prompt,
            code=answer
        )

        db.add(generated_agent)
        db.commit()
        db.refresh(generated_agent)

        saved = save_research(
            user_email=user,
            topic=f"Agent Builder: {agent_type}",
            report=answer,
            logs=[
                "Agent Builder Mode Activated",
                "Custom AI Agent Generated",
                f"AI Model Used: {selected_model}"
            ]
        )

        return {
            "id": saved["id"],
            "agent_id": generated_agent.id,
            "topic": agent_type,
            "answer": answer,
            "report": answer,
            "logs": [
                "Agent Builder Mode Activated",
                "Custom AI Agent Generated",
                f"AI Model Used: {selected_model}"
            ]
        }

    finally:
        db.close()


@app.get("/my-agents")
async def my_agents(user=Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    db = SessionLocal()

    try:
        agents = (
            db.query(GeneratedAgent)
            .filter(GeneratedAgent.user_email == user)
            .order_by(GeneratedAgent.created_at.desc())
            .all()
        )

        return [
            {
                "id": agent.id,
                "agent_name": agent.agent_name,
                "agent_type": agent.agent_type,
                "role": agent.role,
                "goal": agent.goal,
                "code": agent.code,
                "created_at": str(agent.created_at)
            }
            for agent in agents
        ]

    finally:
        db.close()


@app.delete("/my-agents/{agent_id}")
async def delete_agent(
    agent_id: int,
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    db = SessionLocal()

    try:
        agent = (
            db.query(GeneratedAgent)
            .filter(
                GeneratedAgent.id == agent_id,
                GeneratedAgent.user_email == user
            )
            .first()
        )

        if not agent:
            raise HTTPException(status_code=404, detail="Agent not found")

        db.delete(agent)
        db.commit()

        return {"message": "Agent deleted successfully"}

    finally:
        db.close()


@app.post("/run-agent")
async def run_agent(
    request: RunAgentRequest,
    http_request: Request,
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    selected_model = get_selected_model(http_request)

    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Query is required")

    db = SessionLocal()

    try:
        agent = (
            db.query(GeneratedAgent)
            .filter(
                GeneratedAgent.id == request.agent_id,
                GeneratedAgent.user_email == user
            )
            .first()
        )

        if not agent:
            raise HTTPException(status_code=404, detail="Agent not found")

        prompt = f"""
You are now running this custom AI agent.

Selected AI Model:
{selected_model}

Agent Name:
{agent.agent_name}

Agent Type:
{agent.agent_type}

Agent Goal:
{agent.goal}

Agent System Prompt / Blueprint:
{agent.code}

User Query:
{request.query}

Rules:
- Act like the generated agent.
- Give practical, direct, professional answer.
- Follow the agent role and workflow.
"""

        answer = await ask_llm(prompt, selected_model)

        saved = save_research(
            user_email=user,
            topic=f"Run Agent: {agent.agent_name}",
            report=answer,
            logs=[
                "Custom Agent Run Mode Activated",
                f"Agent Used: {agent.agent_name}",
                f"AI Model Used: {selected_model}"
            ]
        )

        return {
            "id": saved["id"],
            "agent_id": agent.id,
            "agent_name": agent.agent_name,
            "answer": answer,
            "report": answer,
            "logs": [
                "Custom Agent Run Mode Activated",
                f"Agent Used: {agent.agent_name}",
                f"AI Model Used: {selected_model}"
            ]
        }

    finally:
        db.close()


@app.post("/seed-marketplace-agents")
async def seed_marketplace_agents():
    db = SessionLocal()

    agents = [
        {
            "name": "Real Estate Agent",
            "category": "Business",
            "icon": "🏠",
            "description": "Helps users find, compare and analyze properties.",
            "system_prompt": "You are a Real Estate AI Agent. Help users find properties, compare prices, analyze budgets, suggest locations and guide buying or renting decisions."
        },
        {
            "name": "Finance Advisor",
            "category": "Finance",
            "icon": "💰",
            "description": "Helps users with budgeting and financial advice.",
            "system_prompt": "You are a Finance Advisor AI Agent. Help users with budgeting, savings, investment basics and financial planning."
        },
        {
            "name": "Legal Assistant",
            "category": "Legal",
            "icon": "⚖️",
            "description": "Explains legal documents in simple language.",
            "system_prompt": "You are a Legal Assistant AI Agent. Explain legal documents, summarize clauses and provide simple guidance. Do not replace a professional lawyer."
        },
        {
            "name": "Medical Assistant",
            "category": "Health",
            "icon": "🩺",
            "description": "Helps analyze health information.",
            "system_prompt": "You are a Medical Assistant AI Agent. Explain medical reports and health information in simple language. Always suggest consulting a doctor."
        },
        {
            "name": "Travel Planner",
            "category": "Travel",
            "icon": "✈️",
            "description": "Plans trips and travel itineraries.",
            "system_prompt": "You are a Travel Planner AI Agent. Create travel plans, itineraries, budget plans and destination suggestions."
        },
        {
            "name": "Edu Mentor",
            "category": "Education",
            "icon": "🎓",
            "description": "Guides students and learners.",
            "system_prompt": "You are an Education Mentor AI Agent. Help students with study plans, course guidance, learning resources and exam preparation."
        }
    ]

    try:
        for item in agents:
            exists = (
                db.query(MarketplaceAgent)
                .filter(MarketplaceAgent.name == item["name"])
                .first()
            )

            if not exists:
                db.add(MarketplaceAgent(**item))

        db.commit()

        return {"message": "Marketplace agents seeded successfully"}

    finally:
        db.close()


@app.get("/marketplace-agents")
async def get_marketplace_agents(user=Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    db = SessionLocal()

    try:
        agents = db.query(MarketplaceAgent).all()

        return [
            {
                "id": agent.id,
                "name": agent.name,
                "category": agent.category,
                "description": agent.description,
                "icon": agent.icon,
                "system_prompt": agent.system_prompt
            }
            for agent in agents
        ]

    finally:
        db.close()


@app.post("/run-marketplace-agent/{agent_id}")
async def run_marketplace_agent(
    agent_id: int,
    data: dict,
    http_request: Request,
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    selected_model = get_selected_model(http_request)

    query = data.get("query", "").strip()

    if not query:
        raise HTTPException(status_code=400, detail="Query is required")

    db = SessionLocal()

    try:
        agent = (
            db.query(MarketplaceAgent)
            .filter(MarketplaceAgent.id == agent_id)
            .first()
        )

        if not agent:
            raise HTTPException(status_code=404, detail="Agent not found")

        prompt = f"""
{agent.system_prompt}

Selected AI Model:
{selected_model}

User Query:
{query}

Rules:
- Act as {agent.name}.
- Give practical and clear answer.
- Keep answer professional.
"""

        answer = await ask_llm(prompt, selected_model)

        saved = save_research(
            user_email=user,
            topic=f"Marketplace Agent: {agent.name}",
            report=answer,
            logs=[
                "Marketplace Agent Mode Activated",
                f"Agent Used: {agent.name}",
                f"AI Model Used: {selected_model}"
            ]
        )

        return {
            "id": saved["id"],
            "agent_name": agent.name,
            "answer": answer,
            "report": answer,
            "logs": [
                "Marketplace Agent Mode Activated",
                f"Agent Used: {agent.name}",
                f"AI Model Used: {selected_model}"
            ]
        }

    finally:
        db.close()


@app.post("/download-docx")
async def download_docx(data: dict):
    report = data.get("report", "")

    if not report:
        raise HTTPException(status_code=400, detail="Report is required")

    filename = f"{uuid.uuid4()}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)

    doc = Document()
    doc.add_heading("NeuroFlow AI Research Report", level=1)

    for line in report.split("\n"):
        clean_line = line.strip()

        if not clean_line:
            continue

        if clean_line.startswith("# "):
            doc.add_heading(clean_line.replace("# ", ""), level=1)
        elif clean_line.startswith("## "):
            doc.add_heading(clean_line.replace("## ", ""), level=2)
        elif clean_line.startswith("### "):
            doc.add_heading(clean_line.replace("### ", ""), level=3)
        elif clean_line.startswith("- "):
            doc.add_paragraph(clean_line.replace("- ", ""), style="List Bullet")
        else:
            doc.add_paragraph(clean_line)

    doc.save(file_path)

    return FileResponse(
        file_path,
        filename="neuroflow_research_report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@app.post("/download-pdf")
async def download_pdf(data: dict):
    report = data.get("report", "")

    if not report:
        raise HTTPException(status_code=400, detail="Report is required")

    filename = f"{uuid.uuid4()}.pdf"
    file_path = os.path.join(DOWNLOAD_DIR, filename)

    pdf = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("NeuroFlow AI Research Report", styles["Title"]))
    story.append(Spacer(1, 12))

    for line in report.split("\n"):
        clean_line = line.strip()

        if not clean_line:
            story.append(Spacer(1, 8))
            continue

        clean_line = (
            clean_line
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

        if clean_line.startswith("# "):
            story.append(Paragraph(clean_line.replace("# ", ""), styles["Heading1"]))
        elif clean_line.startswith("## "):
            story.append(Paragraph(clean_line.replace("## ", ""), styles["Heading2"]))
        elif clean_line.startswith("### "):
            story.append(Paragraph(clean_line.replace("### ", ""), styles["Heading3"]))
        elif clean_line.startswith("- "):
            story.append(Paragraph("• " + clean_line.replace("- ", ""), styles["BodyText"]))
        else:
            story.append(Paragraph(clean_line, styles["BodyText"]))

        story.append(Spacer(1, 6))

    pdf.build(story)

    return FileResponse(
        file_path,
        filename="neuroflow_research_report.pdf",
        media_type="application/pdf"
    )


@app.post("/change-password")
async def change_password(
    request: ChangePasswordRequest,
    user=Depends(get_current_user)
):
    if not user:
        raise HTTPException(status_code=401, detail="Unauthorized")

    db = SessionLocal()

    try:
        db_user = db.query(User).filter(User.email == user).first()

        if not db_user:
            raise HTTPException(status_code=404, detail="User not found")

        if not verify_password(request.old_password, db_user.password):
            raise HTTPException(status_code=400, detail="Old password is incorrect")

        db_user.password = hash_password(request.new_password)
        db.commit()

        return {"message": "Password changed successfully"}

    finally:
        db.close()


@app.get("/test-llm")
async def test_llm(request: Request):
    selected_model = get_selected_model(request)

    result = await ask_llm(
        "Reply in one line: NeuroFlow AI is working.",
        selected_model
    )

    return {
        "result": result,
        "model": selected_model
    }


@app.get("/auth/google/login")
async def google_login(request: Request):
    redirect_uri = os.getenv(
        "GOOGLE_REDIRECT_URI",
        "http://127.0.0.1:8000/auth/google/callback"
    )

    return await oauth.google.authorize_redirect(request, redirect_uri)


@app.get("/auth/google/callback")
async def google_callback(request: Request):
    frontend_url = os.getenv(
        "FRONTEND_URL",
        "https://neuroflow-ai-eta.vercel.app/"
    )

    try:
        token = await oauth.google.authorize_access_token(request)
        user_info = token.get("userinfo")

        if not user_info:
            user_info = await oauth.google.parse_id_token(request, token)

        email = user_info.get("email")

        if not email:
            raise HTTPException(status_code=400, detail="Google email not found")

        create_user(email=email, password="GOOGLE_OAUTH_USER")

        access_token = create_access_token({"sub": email})

        return RedirectResponse(
            url=f"{frontend_url}/oauth-success?token={access_token}"
        )

    except Exception as e:
        print("Google login error:", e)

        return RedirectResponse(
            url=f"{frontend_url}/?error=google_login_failed"
        )